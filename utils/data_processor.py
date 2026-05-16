import pandas as pd
import re
from io import BytesIO


def parse_score_excel(uploaded_file) -> dict:
    """Parse uploaded Excel file with student scores.
    Expected format: 3 sheets - 考试信息, 题目信息, 学生成绩
    Returns dict with keys: exam_info, question_info, scores_df
    """
    xls = pd.ExcelFile(uploaded_file)

    # Read exam info sheet
    exam_info_df = pd.read_excel(xls, sheet_name="考试信息")
    exam_info = dict(zip(exam_info_df["项目"], exam_info_df["内容"]))

    # Read question info sheet
    question_info_df = pd.read_excel(xls, sheet_name="题目信息")

    # Read scores sheet
    scores_df = pd.read_excel(xls, sheet_name="学生成绩")

    return {
        "exam_info": exam_info,
        "question_info": question_info_df,
        "scores_df": scores_df,
    }


def latex_to_text(text: str) -> str:
    """Convert LaTeX math notation in text to readable Unicode plain text."""
    if not text:
        return text

    # Inline LaTeX: \(...\)
    text = re.sub(r'\\\((.+?)\\\)', lambda m: _convert_latex(m.group(1)), text)
    # Display LaTeX: \[...\] or $$...$$
    text = re.sub(r'\\\[(.+?)\\\]', lambda m: _convert_latex(m.group(1)), text)
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: _convert_latex(m.group(1)), text)
    # Single $...$ (not preceded/followed by another $)
    text = re.sub(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', lambda m: _convert_latex(m.group(1)), text)

    return text.strip()


def _convert_latex(latex: str) -> str:
    """Convert a LaTeX math expression to Unicode text."""
    s = latex

    # \frac{a}{b} -> a/b
    for _ in range(10):
        m = re.search(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', s)
        if not m:
            break
        s = s[:m.start()] + f"({m.group(1)}/{m.group(2)})" + s[m.end():]

    # \sqrt{x} -> √x, \sqrt[n]{x} -> ⁿ√x
    for _ in range(10):
        m = re.search(r'\\sqrt(?:\[([^\]]*)\])?\{([^{}]*)\}', s)
        if not m:
            break
        if m.group(1):
            s = s[:m.start()] + f"{m.group(1)}√({m.group(2)})" + s[m.end():]
        else:
            s = s[:m.start()] + f"√({m.group(2)})" + s[m.end():]

    # \text{} -> text
    s = re.sub(r'\\text\{([^{}]*)\}', r'\1', s)
    # \mathrm{} -> text
    s = re.sub(r'\\mathrm\{([^{}]*)\}', r'\1', s)

    # Superscript: ^x and ^{...}
    s = re.sub(r'\^\{([^{}]*)\}', lambda m: _to_superscript(m.group(1)), s)
    s = re.sub(r'\^(\d)', lambda m: _to_superscript(m.group(1)), s)
    # Subscript: _x and _{...}
    s = re.sub(r'_\{([^{}]*)\}', lambda m: _to_subscript(m.group(1)), s)
    s = re.sub(r'_(\d)', lambda m: _to_subscript(m.group(1)), s)

    # Operators and symbols
    replacements = {
        r'\times': '×', r'\cdot': '·', r'\div': '÷',
        r'\pm': '±', r'\mp': '∓', r'\neq': '≠', r'\ne': '≠',
        r'\leq': '≤', r'\le': '≤', r'\geq': '≥', r'\ge': '≥',
        r'\approx': '≈', r'\equiv': '≡',
        r'\infty': '∞', r'\propto': '∝',
        r'\angle': '∠', r'\perp': '⊥', r'\parallel': '∥',
        r'\triangle': '△', r'\square': '□', r'\circ': '°',
        r'\degree': '°', r'\pi': 'π',
        r'\rightarrow': '→', r'\to': '→', r'\leftarrow': '←',
        r'\Rightarrow': '⇒', r'\Leftarrow': '⇐',
        r'\sum': '∑', r'\prod': '∏',
        r'\int': '∫',
        r'\cup': '∪', r'\cap': '∩', r'\in': '∈',
        r'\subset': '⊂', r'\supset': '⊃',
        r'\log': 'log', r'\ln': 'ln',
        r'\sin': 'sin', r'\cos': 'cos', r'\tan': 'tan',
        r'\cot': 'cot', r'\sec': 'sec', r'\csc': 'csc',
        r'\quad': '  ', r'\qquad': '    ',
    }
    for latex_cmd, unicode_char in replacements.items():
        s = s.replace(latex_cmd, unicode_char)

    # \left \right removal (just remove the markers)
    s = re.sub(r'\\left', '', s)
    s = re.sub(r'\\right', '', s)
    # Remaining LaTeX commands: \foo -> foo (strip backslash, keep name)
    s = re.sub(r'\\([a-zA-Z]+)', r'\1', s)

    # Clean up braces
    s = s.replace('{', '').replace('}', '')
    # Clean up extra whitespace
    s = re.sub(r' +', ' ', s)
    s = re.sub(r' +\)', ')', s)
    s = re.sub(r'\( +', '(', s)

    return s


SUPERSCRIPT_MAP = str.maketrans({
    '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
    '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
    '+': '⁺', '-': '⁻', '(': '⁽', ')': '⁾',
    'a': 'ᵃ', 'b': 'ᵇ', 'c': 'ᶜ', 'd': 'ᵈ', 'e': 'ᵉ',
    'n': 'ⁿ', 'x': 'ˣ', 'y': 'ʸ', 'z': 'ᶻ',
})

SUBSCRIPT_MAP = str.maketrans({
    '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
    '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
    '+': '₊', '-': '₋', '(': '₍', ')': '₎',
    'a': 'ₐ', 'e': 'ₑ', 'i': 'ᵢ', 'j': 'ⱼ',
    'k': 'ₖ', 'n': 'ₙ', 'o': 'ₒ', 'r': 'ᵣ',
    's': 'ₛ', 't': 'ₜ', 'x': 'ₓ',
})


def _to_superscript(s: str) -> str:
    return s.translate(SUPERSCRIPT_MAP)


def _to_subscript(s: str) -> str:
    return s.translate(SUBSCRIPT_MAP)


def _set_chinese_font(run, font_name="宋体", size=12):
    """Set Chinese font for a run, ensuring both Western and East Asian fonts are set."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.font.name = font_name
    run.font.size = Pt(size)
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def export_to_word(title: str, content_sections: list, include_answers: bool = True) -> BytesIO:
    """Export content to a Word document.
    content_sections: list of dicts with 'heading' and 'content' keys.
    include_answers: if False, skip answer/explanation fields (student version).
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Set default font for Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    # Set East Asian font in style XML
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        _set_chinese_font(run, "黑体", 22)

    for section in content_sections:
        # Skip answer sections in student version
        if not include_answers and section.get("heading") == "参考答案与解析":
            continue

        if section.get("heading"):
            heading = doc.add_heading(section["heading"], level=1)
            for run in heading.runs:
                _set_chinese_font(run, "黑体", 16)

        if section.get("content"):
            for line in section["content"].split("\n"):
                if line.strip():
                    p = doc.add_paragraph(latex_to_text(line))
                    for run in p.runs:
                        _set_chinese_font(run)

        if section.get("table"):
            table_data = section["table"]
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                # Header row
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = str(h)
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            _set_chinese_font(run, size=10)
                    # Header background color
                    from docx.oxml import OxmlElement
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:val"), "clear")
                    shading.set(qn("w:color"), "auto")
                    shading.set(qn("w:fill"), "4A90D9")
                    cell._element.get_or_add_tcPr().append(shading)
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = None
                        from docx.shared import RGBColor
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Data rows
                for r_idx, row_data in enumerate(rows):
                    for c_idx, val in enumerate(row_data):
                        cell = table.rows[r_idx + 1].cells[c_idx]
                        cell.text = str(val)
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                _set_chinese_font(run, size=10)
                    # Alternate row color
                    if r_idx % 2 == 0:
                        for c_idx in range(len(headers)):
                            shading = OxmlElement("w:shd")
                            shading.set(qn("w:val"), "clear")
                            shading.set(qn("w:color"), "auto")
                            shading.set(qn("w:fill"), "F2F2F2")
                            table.rows[r_idx + 1].cells[c_idx]._element.get_or_add_tcPr().append(shading)
                doc.add_paragraph("")  # spacing after table

        if section.get("questions"):
            for i, q in enumerate(section["questions"], 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. {latex_to_text(q.get('question', ''))}")
                run.bold = True
                _set_chinese_font(run)
                if q.get("options"):
                    for j, opt in enumerate(q["options"]):
                        # 去除AI可能已添加的字母前缀，统一用A/B/C/D
                        clean_opt = re.sub(r'^[A-Fa-f][.、．]\s*', '', opt.strip())
                        label = chr(65 + j) if j < 26 else str(j + 1)
                        op = doc.add_paragraph(f"    {label}. {latex_to_text(clean_opt)}")
                        for r in op.runs:
                            _set_chinese_font(r)
                if include_answers:
                    if q.get("knowledge_point"):
                        p_kp = doc.add_paragraph()
                        r1 = p_kp.add_run("知识点：")
                        r1.bold = True
                        _set_chinese_font(r1)
                        r2 = p_kp.add_run(latex_to_text(q["knowledge_point"]))
                        _set_chinese_font(r2)
                    if q.get("answer"):
                        p2 = doc.add_paragraph()
                        r1 = p2.add_run("答案：")
                        r1.bold = True
                        _set_chinese_font(r1)
                        r2 = p2.add_run(latex_to_text(q["answer"]))
                        _set_chinese_font(r2)
                    if q.get("explanation"):
                        p3 = doc.add_paragraph()
                        r1 = p3.add_run("解析：")
                        r1.bold = True
                        _set_chinese_font(r1)
                        r2 = p3.add_run(latex_to_text(q["explanation"]))
                        _set_chinese_font(r2)
                doc.add_paragraph("")  # blank line between questions

    # AI 免责声明
    doc.add_paragraph("")
    disclaimer_heading = doc.add_heading("AI生成内容声明", level=2)
    for run in disclaimer_heading.runs:
        _set_chinese_font(run, "黑体", 14)
    disclaimer_text = (
        "本试卷题目由人工智能模型（DeepSeek）自动生成，仅供参考。"
        "题目内容可能存在不准确之处，请教师在使用前进行审核和调整。"
    )
    p_disclaimer = doc.add_paragraph(disclaimer_text)
    for run in p_disclaimer.runs:
        _set_chinese_font(run)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _find_chinese_font():
    """Find a Chinese font file on the system."""
    import os
    font_paths = [
        r"C:\Windows\Fonts\msyh.ttc",    # 微软雅黑
        r"C:\Windows\Fonts\simhei.ttf",   # 黑体
        r"C:\Windows\Fonts\simsun.ttc",   # 宋体
        r"C:\Windows\Fonts\msyhbd.ttc",   # 微软雅黑粗体
    ]
    for path in font_paths:
        if os.path.exists(path):
            return path
    return None


def export_to_pdf(title: str, content_sections: list) -> BytesIO:
    """Export content to a PDF document with Chinese support.
    content_sections: list of dicts with 'heading' and 'content' keys.
    """
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            pass

        def footer(self):
            self.set_y(-15)
            self.set_font("chinese", size=8)
            self.cell(0, 10, f"第 {self.page_no()} 页", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Add Chinese font
    font_path = _find_chinese_font()
    if font_path:
        pdf.add_font("chinese", "", font_path, uni=True)
        pdf.add_font("chinese", "B", font_path, uni=True)
    else:
        # Fallback: try to use a font that might work
        pdf.add_font("chinese", "", "msyh.ttc", uni=True)

    pdf.add_page()

    # Title
    pdf.set_font("chinese", "B", 18)
    pdf.cell(0, 15, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    for section in content_sections:
        if section.get("heading"):
            pdf.set_font("chinese", "B", 14)
            pdf.cell(0, 10, section["heading"], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        if section.get("content"):
            pdf.set_font("chinese", "", 11)
            for line in section["content"].split("\n"):
                if line.strip():
                    pdf.multi_cell(0, 7, line)
                    pdf.ln(1)

        if section.get("table"):
            table = section["table"]
            if table.get("headers") and table.get("rows"):
                headers = table["headers"]
                rows = table["rows"]
                col_width = (pdf.w - 2 * pdf.l_margin) / len(headers)

                # Header
                pdf.set_font("chinese", "B", 10)
                pdf.set_fill_color(74, 144, 217)
                pdf.set_text_color(255, 255, 255)
                for h in headers:
                    pdf.cell(col_width, 8, str(h), border=1, fill=True, align="C")
                pdf.ln()

                # Rows
                pdf.set_font("chinese", "", 10)
                pdf.set_text_color(0, 0, 0)
                for i, row in enumerate(rows):
                    if i % 2 == 0:
                        pdf.set_fill_color(242, 242, 242)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    for cell in row:
                        pdf.cell(col_width, 7, str(cell), border=1, fill=True, align="C")
                    pdf.ln()
                pdf.ln(3)

    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer
